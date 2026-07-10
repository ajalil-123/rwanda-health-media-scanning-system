"""
Offline tests for the Rwanda Health Media Scanning System pipeline.

These use synthetic RSS XML and item dicts -- no network access needed --
so the collection/filtering/dedup/scoring/windowing/storage logic can be
verified without hitting live feeds. Run with:

    python -m unittest discover tests
"""

import os
import sys
import tempfile
import unittest
import unittest.mock
from datetime import date, datetime, timedelta, timezone

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import db  # noqa: E402
import config  # noqa: E402
import report_generator  # noqa: E402
from collectors.rss_utils import parse_feed, parse_rss_datetime  # noqa: E402
from collectors import google_news, web_scraper  # noqa: E402
from processing import filter_relevance, dedup, highlight_score  # noqa: E402
import scan  # noqa: E402


SAMPLE_RSS = """<?xml version="1.0" encoding="UTF-8"?>
<rss version="2.0">
<channel>
  <title>Sample Feed</title>
  <item>
    <title>New malaria treatment rolled out in Rwanda hospitals</title>
    <link>https://example.com/malaria-treatment</link>
    <pubDate>Fri, 10 Jul 2026 09:00:00 GMT</pubDate>
    <description>The Ministry of Health announced a new malaria protocol.</description>
  </item>
  <item>
    <title>Kigali football club wins regional tournament</title>
    <link>https://example.com/football</link>
    <pubDate>Fri, 10 Jul 2026 08:00:00 GMT</pubDate>
    <description>A sports story with no health relevance.</description>
  </item>
</channel>
</rss>
"""

# Mirrors real Google News RSS output: title has a " - Publisher" suffix,
# and a separate <source> tag carries the clean publisher name.
SAMPLE_GOOGLE_NEWS_RSS = """<?xml version="1.0" encoding="UTF-8"?>
<rss version="2.0">
<channel>
  <title>Google News Search</title>
  <item>
    <title>Rwanda expands malaria vaccination programme - The New Times</title>
    <link>https://news.google.com/rss/articles/abc123</link>
    <pubDate>Fri, 10 Jul 2026 07:00:00 GMT</pubDate>
    <description>Rwanda expands malaria vaccination programme</description>
    <source url="https://www.newtimes.co.rw">The New Times</source>
  </item>
  <item>
    <title>Rwanda health ministry responds to outbreak - BBC News</title>
    <link>https://news.google.com/rss/articles/def456</link>
    <pubDate>Fri, 10 Jul 2026 06:00:00 GMT</pubDate>
    <description>Rwanda health ministry responds to outbreak</description>
    <source url="https://www.bbc.com">BBC News</source>
  </item>
</channel>
</rss>
"""

SAMPLE_ATOM = """<?xml version="1.0" encoding="UTF-8"?>
<feed xmlns="http://www.w3.org/2005/Atom">
  <title>Sample Atom Feed</title>
  <entry>
    <title>Rwanda Biomedical Centre launches vaccination drive</title>
    <link href="https://example.com/vaccination"/>
    <updated>2026-07-09T12:00:00Z</updated>
    <summary>RBC announced a new vaccination campaign this week.</summary>
  </entry>
</feed>
"""


class TestFeedParsing(unittest.TestCase):
    def test_parses_rss2_items(self):
        items = parse_feed(SAMPLE_RSS)
        self.assertEqual(len(items), 2)
        self.assertEqual(items[0]["title"], "New malaria treatment rolled out in Rwanda hospitals")
        self.assertIsNotNone(items[0]["published_at"])

    def test_parses_atom_entries(self):
        items = parse_feed(SAMPLE_ATOM)
        self.assertEqual(len(items), 1)
        self.assertEqual(items[0]["url"], "https://example.com/vaccination")

    def test_malformed_xml_returns_empty_list_not_crash(self):
        items = parse_feed("<not valid xml")
        self.assertEqual(items, [])

    def test_rfc2822_date_parses(self):
        dt = parse_rss_datetime("Fri, 10 Jul 2026 09:00:00 GMT")
        self.assertEqual(dt.year, 2026)
        self.assertEqual(dt.month, 7)
        self.assertEqual(dt.day, 10)

    def test_iso_date_parses(self):
        dt = parse_rss_datetime("2026-07-09T12:00:00Z")
        self.assertEqual(dt.day, 9)

    def test_unparseable_date_returns_none(self):
        self.assertIsNone(parse_rss_datetime("not a date"))
        self.assertIsNone(parse_rss_datetime(None))

    def test_source_tag_is_extracted(self):
        items = parse_feed(SAMPLE_GOOGLE_NEWS_RSS)
        self.assertEqual(len(items), 2)
        self.assertEqual(items[0]["source_name"], "The New Times")
        self.assertEqual(items[1]["source_name"], "BBC News")

    def test_feed_without_source_tag_has_none_source_name(self):
        items = parse_feed(SAMPLE_RSS)
        self.assertIsNone(items[0]["source_name"])


class TestGoogleNewsCollector(unittest.TestCase):
    def test_strip_source_suffix_removes_publisher_name(self):
        title = google_news._strip_source_suffix(
            "Rwanda expands malaria vaccination programme - The New Times", "The New Times"
        )
        self.assertEqual(title, "Rwanda expands malaria vaccination programme")

    def test_strip_source_suffix_leaves_title_unchanged_if_no_match(self):
        title = google_news._strip_source_suffix("Some headline", "Unrelated Source")
        self.assertEqual(title, "Some headline")

    def test_strip_source_suffix_handles_missing_source(self):
        title = google_news._strip_source_suffix("Some headline", None)
        self.assertEqual(title, "Some headline")

    def test_known_local_outlet_classified_local(self):
        self.assertEqual(google_news._classify_category("The New Times"), "local_online")
        self.assertEqual(google_news._classify_category("IGIHE"), "local_online")

    def test_unknown_outlet_classified_international(self):
        self.assertEqual(google_news._classify_category("BBC News"), "international")

    def test_unattributed_item_defaults_to_local_online(self):
        self.assertEqual(google_news._classify_category(None), "local_online")

    def test_collect_uses_real_publisher_names_not_generic_label(self):
        with unittest.mock.patch("collectors.google_news.fetch_url", return_value=SAMPLE_GOOGLE_NEWS_RSS):
            with unittest.mock.patch("config.GOOGLE_NEWS_QUERIES", [{"query": "health Rwanda", "hl": "en-RW", "gl": "RW", "ceid": "RW:en"}]):
                items = google_news.collect()
        source_names = {item["source_name"] for item in items}
        self.assertEqual(source_names, {"The New Times", "BBC News"})
        self.assertNotIn("Google News", source_names)
        # Category classification should route each outlet correctly
        categories = {item["source_name"]: item["source_category"] for item in items}
        self.assertEqual(categories["The New Times"], "local_online")
        self.assertEqual(categories["BBC News"], "international")


class TestRelevanceFilter(unittest.TestCase):
    def test_health_item_is_relevant(self):
        item = {"title": "New malaria treatment rolled out in Rwanda hospitals", "summary": ""}
        is_rel, matches = filter_relevance.is_relevant(item)
        self.assertTrue(is_rel)
        self.assertIn("malaria", matches)
        self.assertIn("hospital", matches)

    def test_non_health_item_is_not_relevant(self):
        item = {"title": "Kigali football club wins regional tournament", "summary": "A sports story."}
        is_rel, matches = filter_relevance.is_relevant(item)
        self.assertFalse(is_rel)
        self.assertEqual(matches, [])

    def test_kinyarwanda_keyword_matches(self):
        item = {"title": "Ubuzima bw'ababyeyi butera impungenge", "summary": ""}
        is_rel, matches = filter_relevance.is_relevant(item)
        self.assertTrue(is_rel)

    def test_french_keyword_matches(self):
        item = {"title": "Le ministere de la sante annonce une nouvelle campagne", "summary": ""}
        is_rel, matches = filter_relevance.is_relevant(item)
        self.assertTrue(is_rel)

    def test_filter_items_annotates_matched_keywords(self):
        items = [
            {"title": "Malaria cases rise", "summary": "", "source_name": "A"},
            {"title": "Football scores", "summary": "", "source_name": "B"},
        ]
        result = filter_relevance.filter_items(items)
        self.assertEqual(len(result), 1)
        self.assertIn("matched_keywords", result[0])

    def test_filter_does_not_mutate_input(self):
        items = [{"title": "Malaria cases rise", "summary": "", "source_name": "A"}]
        filter_relevance.filter_items(items)
        self.assertNotIn("matched_keywords", items[0])


class TestDedup(unittest.TestCase):
    def test_near_duplicate_titles_are_folded(self):
        items = [
            {"title": "New malaria treatment rolled out in Rwanda hospitals", "source_name": "Outlet A", "matched_keywords": ["malaria"]},
            {"title": "New malaria treatment rolled out in Rwanda's hospitals", "source_name": "Outlet B", "matched_keywords": ["malaria"]},
        ]
        unique_items, duplicates = dedup.deduplicate(items)
        self.assertEqual(len(unique_items), 1)
        self.assertEqual(len(duplicates), 1)
        self.assertIn("Outlet B", unique_items[0]["covered_by"])

    def test_distinct_stories_are_not_folded(self):
        items = [
            {"title": "New malaria treatment rolled out in hospitals", "source_name": "Outlet A", "matched_keywords": []},
            {"title": "Cholera outbreak reported in Western Province", "source_name": "Outlet B", "matched_keywords": []},
        ]
        unique_items, duplicates = dedup.deduplicate(items)
        self.assertEqual(len(unique_items), 2)
        self.assertEqual(len(duplicates), 0)


class TestHighlightScoring(unittest.TestCase):
    def test_more_outlets_scores_higher(self):
        now = datetime(2026, 7, 10, tzinfo=timezone.utc)
        item_wide = {"source_name": "A", "covered_by": ["A", "B", "C"], "matched_keywords": ["malaria"], "published_at": now}
        item_single = {"source_name": "D", "covered_by": ["D"], "matched_keywords": ["malaria"], "published_at": now}
        self.assertGreater(highlight_score.score_item(item_wide, now=now), highlight_score.score_item(item_single, now=now))

    def test_older_item_scores_lower(self):
        now = datetime(2026, 7, 10, tzinfo=timezone.utc)
        fresh = {"source_name": "A", "matched_keywords": [], "published_at": now}
        old = {"source_name": "B", "matched_keywords": [], "published_at": now - timedelta(days=6)}
        self.assertGreater(highlight_score.score_item(fresh, now=now), highlight_score.score_item(old, now=now))

    def test_rank_items_sorted_descending(self):
        now = datetime(2026, 7, 10, tzinfo=timezone.utc)
        items = [
            {"source_name": "A", "matched_keywords": [], "published_at": now - timedelta(days=5)},
            {"source_name": "B", "covered_by": ["B", "C"], "matched_keywords": ["malaria"], "published_at": now},
        ]
        ranked = highlight_score.rank_items(items, now=now)
        self.assertEqual(ranked[0]["source_name"], "B")


class TestScanWindow(unittest.TestCase):
    def test_daily_window_is_exactly_one_calendar_day(self):
        target = date(2026, 7, 10)
        start, end = scan.compute_window("daily", target_date=target)
        self.assertEqual(start, datetime(2026, 7, 10, 0, 0, 0, tzinfo=timezone.utc))
        self.assertEqual(end, datetime(2026, 7, 10, 23, 59, 59, tzinfo=timezone.utc))

    def test_weekly_window_is_seven_calendar_days_ending_on_target(self):
        target = date(2026, 7, 10)
        start, end = scan.compute_window("weekly", target_date=target)
        self.assertEqual(start, datetime(2026, 7, 4, 0, 0, 0, tzinfo=timezone.utc))
        self.assertEqual(end, datetime(2026, 7, 10, 23, 59, 59, tzinfo=timezone.utc))
        self.assertEqual((end.date() - start.date()).days, 6)  # 7 days inclusive of both ends

    def test_weekly_custom_period_uses_explicit_start_and_end(self):
        """A user-specified period should be used exactly as given, not
        forced to 7 days -- this is what lets someone pick e.g. a 10-day
        period covering a specific outbreak response window."""
        start_date = date(2026, 6, 20)
        end_date = date(2026, 6, 30)
        start, end = scan.compute_window("weekly", target_date=end_date, start_date=start_date)
        self.assertEqual(start, datetime(2026, 6, 20, 0, 0, 0, tzinfo=timezone.utc))
        self.assertEqual(end, datetime(2026, 6, 30, 23, 59, 59, tzinfo=timezone.utc))
        self.assertEqual((end.date() - start.date()).days, 10)

    def test_weekly_custom_period_can_be_shorter_than_default(self):
        start_date = date(2026, 7, 8)
        end_date = date(2026, 7, 10)
        start, end = scan.compute_window("weekly", target_date=end_date, start_date=start_date)
        self.assertEqual((end.date() - start.date()).days, 2)

    def test_weekly_start_date_after_target_date_raises(self):
        with self.assertRaises(ValueError):
            scan.compute_window("weekly", target_date=date(2026, 7, 1), start_date=date(2026, 7, 10))

    def test_daily_mode_ignores_start_date(self):
        """start_date is a weekly-only concept -- passing it for a daily
        scan should have no effect."""
        start, end = scan.compute_window("daily", target_date=date(2026, 7, 10), start_date=date(2026, 6, 1))
        self.assertEqual(start.date(), date(2026, 7, 10))
        self.assertEqual(end.date(), date(2026, 7, 10))

    def test_specific_past_date_is_respected_for_daily(self):
        target = date(2026, 6, 15)
        start, end = scan.compute_window("daily", target_date=target)
        self.assertEqual(start.date(), date(2026, 6, 15))
        self.assertEqual(end.date(), date(2026, 6, 15))

    def test_repeated_calls_with_same_date_are_identical(self):
        """Running the same --mode/--date combination must always compute
        the same window, regardless of what time it's actually run at --
        this is the whole point of letting the user specify a date."""
        target = date(2026, 7, 6)
        first = scan.compute_window("weekly", target_date=target)
        second = scan.compute_window("weekly", target_date=target)
        self.assertEqual(first, second)

    def test_datetime_input_is_reduced_to_its_date(self):
        """Passing a full datetime (not just a date) should still work --
        only the date portion should affect the window, so scanning at
        9am vs 11pm on the same requested day gives the same result."""
        morning = datetime(2026, 7, 10, 9, 0, 0, tzinfo=timezone.utc)
        evening = datetime(2026, 7, 10, 23, 0, 0, tzinfo=timezone.utc)
        self.assertEqual(scan.compute_window("daily", target_date=morning),
                          scan.compute_window("daily", target_date=evening))

    def test_no_date_given_defaults_to_today(self):
        start, end = scan.compute_window("daily", target_date=None)
        today = datetime.now(timezone.utc).date()
        self.assertEqual(start.date(), today)
        self.assertEqual(end.date(), today)

    def test_item_inside_window_is_kept(self):
        start = datetime(2026, 7, 3, tzinfo=timezone.utc)
        end = datetime(2026, 7, 10, 23, 59, 59, tzinfo=timezone.utc)
        item = {"published_at": datetime(2026, 7, 5, tzinfo=timezone.utc)}
        self.assertTrue(scan.within_window(item, start, end))

    def test_item_on_the_exact_boundary_is_kept(self):
        start = datetime(2026, 7, 3, 0, 0, 0, tzinfo=timezone.utc)
        end = datetime(2026, 7, 10, 23, 59, 59, tzinfo=timezone.utc)
        item = {"published_at": end}  # exactly at the inclusive upper bound
        self.assertTrue(scan.within_window(item, start, end))

    def test_item_outside_window_is_dropped(self):
        start = datetime(2026, 7, 3, tzinfo=timezone.utc)
        end = datetime(2026, 7, 10, 23, 59, 59, tzinfo=timezone.utc)
        item = {"published_at": datetime(2026, 6, 1, tzinfo=timezone.utc)}
        self.assertFalse(scan.within_window(item, start, end))

    def test_item_with_unknown_date_is_kept_not_dropped(self):
        start = datetime(2026, 7, 3, tzinfo=timezone.utc)
        end = datetime(2026, 7, 10, 23, 59, 59, tzinfo=timezone.utc)
        item = {"published_at": None}
        self.assertTrue(scan.within_window(item, start, end))


class TestDatabase(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        self.tmp.close()
        self.db_path = self.tmp.name
        db.init_db(self.db_path)

    def tearDown(self):
        os.unlink(self.db_path)

    def test_insert_and_retrieve_item(self):
        scan_id = db.start_scan("daily", "2026-07-09T00:00:00", "2026-07-10T00:00:00", "2026-07-10T00:00:00", db_path=self.db_path)
        item = {
            "source_name": "The New Times", "source_category": "local_online",
            "title": "Test story", "url": "https://example.com/1",
            "language": "en", "published_at": "2026-07-10T00:00:00",
            "fetched_at": "2026-07-10T00:00:00", "matched_keywords": ["health"], "snippet": "",
        }
        item_id = db.insert_item(item, scan_id, db_path=self.db_path)
        self.assertIsNotNone(item_id)

        items = db.get_items_for_scan(scan_id, db_path=self.db_path)
        self.assertEqual(len(items), 1)
        self.assertEqual(items[0]["title"], "Test story")

    def test_duplicate_url_is_not_inserted_twice(self):
        scan_id = db.start_scan("daily", "2026-07-09T00:00:00", "2026-07-10T00:00:00", "2026-07-10T00:00:00", db_path=self.db_path)
        item = {
            "source_name": "A", "source_category": "local_online", "title": "T",
            "url": "https://example.com/dup", "language": "en", "published_at": None,
            "fetched_at": "2026-07-10T00:00:00", "matched_keywords": [], "snippet": "",
        }
        first_id = db.insert_item(item, scan_id, db_path=self.db_path)
        second_id = db.insert_item(item, scan_id, db_path=self.db_path)
        self.assertIsNotNone(first_id)
        self.assertIsNone(second_id)

    def test_finish_scan_records_counts(self):
        scan_id = db.start_scan("weekly", "2026-07-03T00:00:00", "2026-07-10T00:00:00", "2026-07-10T00:00:00", db_path=self.db_path)
        db.finish_scan(scan_id, "2026-07-10T00:05:00", raw_items=50, relevant_items=12, unique_items=9, db_path=self.db_path)
        row = db.get_scan(scan_id, db_path=self.db_path)
        self.assertEqual(row["raw_items_collected"], 50)
        self.assertEqual(row["unique_items"], 9)

    def test_same_url_in_two_different_scans_is_kept_in_both(self):
        """The core bug fix: an article whose URL was already stored under
        one scan must still be storable -- and visible -- under a
        different scan. Previously the global UNIQUE(url) constraint
        silently dropped it from the second scan entirely."""
        scan_1 = db.start_scan("daily", "2026-07-09T00:00:00", "2026-07-09T23:59:59", "2026-07-09T00:00:00", db_path=self.db_path)
        scan_2 = db.start_scan("daily", "2026-07-10T00:00:00", "2026-07-10T23:59:59", "2026-07-10T00:00:00", db_path=self.db_path)

        item = {
            "source_name": "The New Times", "source_category": "local_online",
            "title": "A story that stays relevant across two days",
            "url": "https://example.com/persistent-story",
            "language": "en", "published_at": "2026-07-09T10:00:00",
            "fetched_at": "2026-07-09T10:00:00", "matched_keywords": ["health"], "snippet": "",
        }
        first_id = db.insert_item(item, scan_1, db_path=self.db_path)
        second_id = db.insert_item(item, scan_2, db_path=self.db_path)

        self.assertIsNotNone(first_id, "item should be stored for the first scan")
        self.assertIsNotNone(second_id, "the SAME url must also be stored for a different scan -- this was the bug")

        self.assertEqual(len(db.get_items_for_scan(scan_1, db_path=self.db_path)), 1)
        self.assertEqual(len(db.get_items_for_scan(scan_2, db_path=self.db_path)), 1)

    def test_same_url_twice_within_the_same_scan_is_still_deduplicated(self):
        """The per-scan uniqueness should still hold -- e.g. if Google News
        and a direct RSS feed both surface the identical URL in one scan,
        it shouldn't be stored twice under that scan."""
        scan_id = db.start_scan("daily", "2026-07-10T00:00:00", "2026-07-10T23:59:59", "2026-07-10T00:00:00", db_path=self.db_path)
        item = {
            "source_name": "The New Times", "source_category": "local_online", "title": "T",
            "url": "https://example.com/same-scan-dup", "language": "en", "published_at": None,
            "fetched_at": "2026-07-10T00:00:00", "matched_keywords": [], "snippet": "",
        }
        first_id = db.insert_item(item, scan_id, db_path=self.db_path)
        second_id = db.insert_item(item, scan_id, db_path=self.db_path)
        self.assertIsNotNone(first_id)
        self.assertIsNone(second_id)
        self.assertEqual(len(db.get_items_for_scan(scan_id, db_path=self.db_path)), 1)

    def test_delete_scan_removes_scan_and_its_items(self):
        scan_id = db.start_scan("daily", "2026-07-10T00:00:00", "2026-07-10T23:59:59", "2026-07-10T00:00:00", db_path=self.db_path)
        db.insert_item({
            "source_name": "A", "source_category": "local_online", "title": "T",
            "url": "https://example.com/to-be-deleted", "language": "en", "published_at": None,
            "fetched_at": "2026-07-10T00:00:00", "matched_keywords": [], "snippet": "",
        }, scan_id, db_path=self.db_path)

        deleted = db.delete_scan(scan_id, db_path=self.db_path)
        self.assertTrue(deleted)
        self.assertIsNone(db.get_scan(scan_id, db_path=self.db_path))
        self.assertEqual(db.get_items_for_scan(scan_id, db_path=self.db_path), [])

    def test_delete_scan_returns_false_for_nonexistent_scan(self):
        self.assertFalse(db.delete_scan(99999, db_path=self.db_path))

    def test_delete_scan_does_not_affect_other_scans(self):
        scan_1 = db.start_scan("daily", "2026-07-09T00:00:00", "2026-07-09T23:59:59", "2026-07-09T00:00:00", db_path=self.db_path)
        scan_2 = db.start_scan("daily", "2026-07-10T00:00:00", "2026-07-10T23:59:59", "2026-07-10T00:00:00", db_path=self.db_path)
        db.insert_item({
            "source_name": "A", "source_category": "local_online", "title": "T1",
            "url": "https://example.com/scan1-item", "language": "en", "published_at": None,
            "fetched_at": "2026-07-09T00:00:00", "matched_keywords": [], "snippet": "",
        }, scan_1, db_path=self.db_path)
        db.insert_item({
            "source_name": "A", "source_category": "local_online", "title": "T2",
            "url": "https://example.com/scan2-item", "language": "en", "published_at": None,
            "fetched_at": "2026-07-10T00:00:00", "matched_keywords": [], "snippet": "",
        }, scan_2, db_path=self.db_path)

        db.delete_scan(scan_1, db_path=self.db_path)

        self.assertIsNone(db.get_scan(scan_1, db_path=self.db_path))
        self.assertIsNotNone(db.get_scan(scan_2, db_path=self.db_path))
        self.assertEqual(len(db.get_items_for_scan(scan_2, db_path=self.db_path)), 1)

    def test_delete_all_scans_clears_everything_and_returns_count(self):
        for i in range(3):
            scan_id = db.start_scan("daily", "2026-07-10T00:00:00", "2026-07-10T23:59:59", "2026-07-10T00:00:00", db_path=self.db_path)
            db.insert_item({
                "source_name": "A", "source_category": "local_online", "title": f"T{i}",
                "url": f"https://example.com/item-{i}", "language": "en", "published_at": None,
                "fetched_at": "2026-07-10T00:00:00", "matched_keywords": [], "snippet": "",
            }, scan_id, db_path=self.db_path)

        deleted_count = db.delete_all_scans(db_path=self.db_path)
        self.assertEqual(deleted_count, 3)
        self.assertEqual(db.list_scans(db_path=self.db_path), [])

    def test_delete_all_scans_on_empty_database_returns_zero(self):
        self.assertEqual(db.delete_all_scans(db_path=self.db_path), 0)

    def test_migration_fixes_a_pre_existing_old_schema_database(self):
        """Simulates exactly the database shape a user already running an
        earlier version of this system would have: UNIQUE(url) only, none
        of the later-added columns. init_db() should detect this, migrate
        it in place, preserve existing data, and the bug should then be
        fixed without the user needing to delete their database."""
        old_schema_db = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        old_schema_db.close()
        try:
            with db.get_connection(old_schema_db.name) as conn:
                conn.executescript("""
                    CREATE TABLE items (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        source_name TEXT NOT NULL,
                        source_category TEXT NOT NULL,
                        title TEXT NOT NULL,
                        url TEXT NOT NULL,
                        language TEXT,
                        published_at TEXT,
                        fetched_at TEXT NOT NULL,
                        matched_keywords TEXT,
                        snippet TEXT,
                        duplicate_of INTEGER,
                        scan_id INTEGER NOT NULL,
                        UNIQUE(url)
                    );
                    CREATE TABLE scans (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        mode TEXT NOT NULL,
                        window_start TEXT NOT NULL,
                        window_end TEXT NOT NULL,
                        started_at TEXT NOT NULL,
                        finished_at TEXT,
                        raw_items_collected INTEGER DEFAULT 0,
                        relevant_items INTEGER DEFAULT 0,
                        unique_items INTEGER DEFAULT 0
                    );
                """)
                conn.execute(
                    "INSERT INTO scans (id, mode, window_start, window_end, started_at) VALUES (1, 'daily', 'x', 'y', 'z')"
                )
                conn.execute(
                    """INSERT INTO items (source_name, source_category, title, url, fetched_at, scan_id)
                       VALUES ('Old Source', 'local_online', 'A pre-existing story', 'https://example.com/pre-existing', 'z', 1)"""
                )

            # This is the moment a user would restart the app after pulling the fix.
            db.init_db(old_schema_db.name)

            # Old data should have survived the migration.
            preserved = db.get_items_for_scan(1, db_path=old_schema_db.name)
            self.assertEqual(len(preserved), 1)
            self.assertEqual(preserved[0]["title"], "A pre-existing story")

            # And the actual bug should now be fixed: the same URL can be
            # stored again under a brand new scan.
            new_scan_id = db.start_scan("daily", "2026-07-10T00:00:00", "2026-07-10T23:59:59", "2026-07-10T00:00:00", db_path=old_schema_db.name)
            new_id = db.insert_item({
                "source_name": "Old Source", "source_category": "local_online",
                "title": "A pre-existing story", "url": "https://example.com/pre-existing",
                "language": "en", "published_at": None, "fetched_at": "2026-07-10T00:00:00",
                "matched_keywords": [], "snippet": "",
            }, new_scan_id, db_path=old_schema_db.name)
            self.assertIsNotNone(new_id, "same URL must now be storable under a new scan after migration")

            # And re-running init_db again should be a safe no-op (idempotent).
            db.init_db(old_schema_db.name)
            self.assertEqual(len(db.get_items_for_scan(1, db_path=old_schema_db.name)), 1)
            self.assertEqual(len(db.get_items_for_scan(new_scan_id, db_path=old_schema_db.name)), 1)
        finally:
            os.unlink(old_schema_db.name)


class TestReportGenerator(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        self.tmp.close()
        self.db_path = self.tmp.name
        self.out_dir = tempfile.mkdtemp()
        db.init_db(self.db_path)
        self.scan_id = db.start_scan(
            "daily", "2026-07-09T00:00:00+00:00", "2026-07-10T00:00:00+00:00",
            "2026-07-10T00:00:00+00:00", db_path=self.db_path,
        )

    def tearDown(self):
        os.unlink(self.db_path)

    def _insert_sample_item(self, category="local_online", source="The New Times", title="Sample story"):
        item = {
            "source_name": source, "source_category": category, "title": title,
            "url": f"https://example.com/{title.replace(' ', '-')}", "language": "en",
            "published_at": "2026-07-10T00:00:00+00:00", "fetched_at": "2026-07-10T00:00:00+00:00",
            "matched_keywords": ["health"], "snippet": "", "highlight_score": 5.0, "covered_by": [source],
        }
        return db.insert_item(item, self.scan_id, db_path=self.db_path)

    def test_raises_if_scan_does_not_exist(self):
        with self.assertRaises(ValueError):
            report_generator.generate_report(99999, db_path=self.db_path, output_dir=self.out_dir)

    def test_raises_if_nothing_included_yet(self):
        self._insert_sample_item()
        # Not marked included -- should refuse rather than generate an empty report
        with self.assertRaises(ValueError):
            report_generator.generate_report(self.scan_id, db_path=self.db_path, output_dir=self.out_dir)

    def test_generates_docx_with_included_items_only(self):
        included_id = self._insert_sample_item(title="Included story")
        excluded_id = self._insert_sample_item(title="Excluded story")
        db.update_item_review(included_id, included=1, editor_summary="A summary.", db_path=self.db_path)
        db.update_item_review(excluded_id, included=0, editor_summary="", db_path=self.db_path)

        out_path = report_generator.generate_report(self.scan_id, db_path=self.db_path, output_dir=self.out_dir)
        self.assertTrue(os.path.exists(out_path))

        from docx import Document
        doc = Document(out_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("Included story", full_text)
        self.assertNotIn("Excluded story", full_text)
        self.assertIn("A summary.", full_text)

    def test_report_contains_all_four_sections(self):
        item_id = self._insert_sample_item()
        db.update_item_review(item_id, included=1, editor_summary="Summary.", db_path=self.db_path)
        out_path = report_generator.generate_report(self.scan_id, db_path=self.db_path, output_dir=self.out_dir)

        from docx import Document
        doc = Document(out_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("I. Local Media", full_text)
        self.assertIn("A. Online", full_text)
        self.assertIn("B. Radio & TV", full_text)
        self.assertIn("II. International Media", full_text)
        self.assertIn("III. Research Findings", full_text)
        self.assertIn("IV. Social Media Trends", full_text)

    def test_hyperlinks_point_to_correct_urls(self):
        item_id = self._insert_sample_item(title="Linked story")
        db.update_item_review(item_id, included=1, editor_summary="Summary.", db_path=self.db_path)
        out_path = report_generator.generate_report(self.scan_id, db_path=self.db_path, output_dir=self.out_dir)

        import zipfile
        with zipfile.ZipFile(out_path) as z:
            rels = z.read("word/_rels/document.xml.rels").decode("utf-8")
        self.assertIn("https://example.com/Linked-story", rels)

    def test_items_without_matching_category_leave_placeholder(self):
        # No research items inserted at all
        item_id = self._insert_sample_item(category="local_online")
        db.update_item_review(item_id, included=1, editor_summary="Summary.", db_path=self.db_path)
        out_path = report_generator.generate_report(self.scan_id, db_path=self.db_path, output_dir=self.out_dir)

        from docx import Document
        doc = Document(out_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        self.assertIn("No research items were marked included", full_text)


class TestScanDiagnostics(unittest.TestCase):
    """Covers the diagnostics added to help distinguish 'nothing was
    collected' (network/access problem) from 'items were collected but
    filtered out' (date window or keyword list) -- the ambiguity that
    made 'no health news is showing up' hard to debug."""

    def setUp(self):
        self.tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        self.tmp.close()
        self.db_path = self.tmp.name

    def tearDown(self):
        os.unlink(self.db_path)

    def test_diagnostics_report_zero_when_all_collectors_return_nothing(self):
        with unittest.mock.patch("collectors.google_news.collect", return_value=[]), \
             unittest.mock.patch("collectors.direct_rss.collect", return_value=[]), \
             unittest.mock.patch("collectors.web_scraper.collect", return_value=[]), \
             unittest.mock.patch("collectors.pubmed.collect", return_value=[]):
            result = scan.run_scan("daily", target_date=date(2026, 7, 10), db_path=self.db_path)
        diag = result["diagnostics"]
        self.assertEqual(diag["raw_items"], 0)
        self.assertEqual(diag["collector_counts"], {"google_news": 0, "direct_rss": 0, "web_scraper": 0, "pubmed": 0})

    def test_diagnostics_distinguish_windowed_out_from_never_collected(self):
        """Items collected but all outside the date window should show up
        as raw_items > 0 but windowed_items == 0 -- not indistinguishable
        from a total collection failure."""
        old_item = {
            "title": "Malaria treatment rolled out", "url": "https://example.com/x",
            "published_at": datetime(2026, 1, 1, tzinfo=timezone.utc), "summary": "",
            "source_name": "The New Times", "source_category": "local_online", "language": "en",
        }
        with unittest.mock.patch("collectors.google_news.collect", return_value=[old_item]), \
             unittest.mock.patch("collectors.direct_rss.collect", return_value=[]), \
             unittest.mock.patch("collectors.web_scraper.collect", return_value=[]), \
             unittest.mock.patch("collectors.pubmed.collect", return_value=[]):
            result = scan.run_scan("daily", target_date=date(2026, 7, 10), db_path=self.db_path)
        diag = result["diagnostics"]
        self.assertEqual(diag["raw_items"], 1)
        self.assertEqual(diag["windowed_items"], 0)

    def test_diagnostics_distinguish_filtered_out_from_never_collected(self):
        """Items in-window but not health-related should show up as
        windowed_items > 0 but relevant_items == 0."""
        football_item = {
            "title": "Rwanda wins football match", "url": "https://example.com/y",
            "published_at": datetime(2026, 7, 10, 10, 0, tzinfo=timezone.utc), "summary": "",
            "source_name": "The New Times", "source_category": "local_online", "language": "en",
        }
        with unittest.mock.patch("collectors.google_news.collect", return_value=[football_item]), \
             unittest.mock.patch("collectors.direct_rss.collect", return_value=[]), \
             unittest.mock.patch("collectors.web_scraper.collect", return_value=[]), \
             unittest.mock.patch("collectors.pubmed.collect", return_value=[]):
            result = scan.run_scan("daily", target_date=date(2026, 7, 10), db_path=self.db_path)
        diag = result["diagnostics"]
        self.assertEqual(diag["windowed_items"], 1)
        self.assertEqual(diag["relevant_items"], 0)

    def test_diagnostics_report_category_breakdown(self):
        items = [
            {"title": "Malaria in hospitals", "url": "https://example.com/a",
             "published_at": datetime(2026, 7, 10, 10, 0, tzinfo=timezone.utc), "summary": "",
             "source_name": "The New Times", "source_category": "local_online", "language": "en"},
            {"title": "WHO reports on malaria", "url": "https://example.com/b",
             "published_at": datetime(2026, 7, 10, 10, 0, tzinfo=timezone.utc), "summary": "",
             "source_name": "BBC", "source_category": "international", "language": "en"},
        ]
        with unittest.mock.patch("collectors.google_news.collect", return_value=items), \
             unittest.mock.patch("collectors.direct_rss.collect", return_value=[]), \
             unittest.mock.patch("collectors.web_scraper.collect", return_value=[]), \
             unittest.mock.patch("collectors.pubmed.collect", return_value=[]):
            result = scan.run_scan("daily", target_date=date(2026, 7, 10), db_path=self.db_path)
        diag = result["diagnostics"]
        self.assertEqual(diag["category_counts"], {"local_online": 1, "international": 1})

    def test_a_crashing_collector_does_not_crash_the_whole_scan(self):
        """If a collector raises (not just returns []), the scan should
        still complete using whatever the other collectors returned."""
        def broken_collector():
            raise RuntimeError("simulated collector crash")

        with unittest.mock.patch("collectors.google_news.collect", side_effect=broken_collector), \
             unittest.mock.patch("collectors.direct_rss.collect", return_value=[]), \
             unittest.mock.patch("collectors.web_scraper.collect", return_value=[]), \
             unittest.mock.patch("collectors.pubmed.collect", return_value=[]):
            result = scan.run_scan("daily", target_date=date(2026, 7, 10), db_path=self.db_path)
        self.assertEqual(result["diagnostics"]["collector_counts"]["google_news"], 0)


class TestWebAppDeleteRoutes(unittest.TestCase):
    """Tests the actual Flask routes for deleting scans, not just the db
    functions -- covers the confirmation-free POST endpoint behavior,
    flash messages, and redirects a browser would actually hit."""

    def setUp(self):
        self.tmp = tempfile.NamedTemporaryFile(suffix=".db", delete=False)
        self.tmp.close()
        self.db_path = self.tmp.name
        db.init_db(self.db_path)

        webapp_dir = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "webapp")
        if webapp_dir not in sys.path:
            sys.path.insert(0, webapp_dir)
        import app as webapp_app  # noqa: PLC0415 -- deliberately imported here, needs webapp_dir on sys.path first
        self.webapp_app = webapp_app
        self.client = webapp_app.app.test_client()

        # Point config.DB_PATH at our temp db so the routes (which call
        # db.init_db() with no path, i.e. the default) operate on it.
        import config
        self._original_db_path = config.DB_PATH
        config.DB_PATH = self.db_path

    def tearDown(self):
        import config
        config.DB_PATH = self._original_db_path
        os.unlink(self.db_path)

    def _make_scan_with_item(self):
        scan_id = db.start_scan("daily", "2026-07-10T00:00:00", "2026-07-10T23:59:59", "2026-07-10T00:00:00", db_path=self.db_path)
        db.insert_item({
            "source_name": "A", "source_category": "local_online", "title": "T",
            "url": f"https://example.com/item-{scan_id}", "language": "en", "published_at": None,
            "fetched_at": "2026-07-10T00:00:00", "matched_keywords": [], "snippet": "",
        }, scan_id, db_path=self.db_path)
        return scan_id

    def test_delete_scan_route_removes_it_and_redirects_with_flash(self):
        scan_id = self._make_scan_with_item()
        resp = self.client.post(f"/scan/{scan_id}/delete", follow_redirects=True)
        self.assertEqual(resp.status_code, 200)
        self.assertIn(b"deleted", resp.data)
        self.assertIsNone(db.get_scan(scan_id, db_path=self.db_path))

    def test_delete_scan_route_for_nonexistent_id_does_not_crash(self):
        resp = self.client.post("/scan/99999/delete", follow_redirects=True)
        self.assertEqual(resp.status_code, 200)
        self.assertIn(b"No scan found", resp.data)

    def test_delete_all_scans_route_clears_everything(self):
        self._make_scan_with_item()
        self._make_scan_with_item()
        resp = self.client.post("/delete-all-scans", follow_redirects=True)
        self.assertEqual(resp.status_code, 200)
        self.assertIn(b"Deleted all 2", resp.data)
        self.assertEqual(db.list_scans(db_path=self.db_path), [])

    def test_index_page_shows_delete_buttons_when_scans_exist(self):
        self._make_scan_with_item()
        resp = self.client.get("/")
        self.assertIn(b"Delete", resp.data)
        self.assertIn(b"Delete all scans", resp.data)

    def test_index_page_has_no_delete_all_button_when_empty(self):
        resp = self.client.get("/")
        self.assertNotIn(b"Delete all scans", resp.data)


class TestWebScraper(unittest.TestCase):
    """Tests the generic scraping heuristic and configured-selector path
    against synthetic HTML, since this sandbox has no live network access
    to verify against real site structures."""

    WORDPRESS_STYLE_HTML = """
    <html><body>
      <nav>
        <a href="/">Home</a>
        <a href="/about">About Us</a>
        <a href="/contact">Contact</a>
      </nav>
      <article>
        <h2 class="entry-title"><a href="/2026/07/malaria-vaccination-drive-launched-in-rwanda">Malaria vaccination drive launched in Rwanda</a></h2>
      </article>
      <article>
        <h2 class="entry-title"><a href="/2026/07/health-ministry-responds-to-outbreak-concerns">Health ministry responds to outbreak concerns</a></h2>
      </article>
      <footer>
        <a href="/privacy">Privacy Policy</a>
        <a href="/wp-login.php">Login</a>
      </footer>
    </body></html>
    """

    GENERIC_HTML = """
    <html><body>
      <div class="menu">
        <a href="/">Home</a>
        <a href="/sports">Sports</a>
      </div>
      <div class="story">
        <a href="/news/rwanda-launches-new-hospital-in-musanze-district-2026">Rwanda launches new hospital in Musanze district</a>
      </div>
      <div class="story">
        <a href="https://external-ads.example.com/click?id=123">Buy now - amazing deals</a>
      </div>
      <div class="story">
        <a href="/news/x">Hi</a>
      </div>
    </body></html>
    """

    def test_configured_selector_extracts_only_matching_links(self):
        soup = __import__("bs4").BeautifulSoup(self.WORDPRESS_STYLE_HTML, "html.parser")
        found = web_scraper._extract_with_selector(soup, "h2.entry-title a", "https://example.com/")
        self.assertEqual(len(found), 2)
        urls = [f[0] for f in found]
        self.assertIn("https://example.com/2026/07/malaria-vaccination-drive-launched-in-rwanda", urls)

    def test_configured_selector_ignores_nav_and_footer_links(self):
        soup = __import__("bs4").BeautifulSoup(self.WORDPRESS_STYLE_HTML, "html.parser")
        found = web_scraper._extract_with_selector(soup, "h2.entry-title a", "https://example.com/")
        titles = [f[1] for f in found]
        self.assertNotIn("Home", titles)
        self.assertNotIn("Contact", titles)
        self.assertNotIn("Login", titles)

    def test_generic_heuristic_finds_headline_via_common_selector(self):
        soup = __import__("bs4").BeautifulSoup(self.WORDPRESS_STYLE_HTML, "html.parser")
        found = web_scraper._extract_generic(soup, "https://example.com/", "example.com")
        self.assertEqual(len(found), 2)

    def test_generic_heuristic_filters_short_and_offsite_links(self):
        soup = __import__("bs4").BeautifulSoup(self.GENERIC_HTML, "html.parser")
        found = web_scraper._extract_generic(soup, "https://example.com/", "example.com")
        titles = [f[1] for f in found]
        # "Hi" is too short, the ad link is off-site -- neither should survive
        self.assertNotIn("Hi", titles)
        self.assertFalse(any("Buy now" in t for t in titles))
        self.assertTrue(any("Musanze" in t for t in titles))

    def test_relative_urls_are_resolved_to_absolute(self):
        soup = __import__("bs4").BeautifulSoup(self.WORDPRESS_STYLE_HTML, "html.parser")
        found = web_scraper._extract_with_selector(soup, "h2.entry-title a", "https://example.com/news/")
        for url, _ in found:
            self.assertTrue(url.startswith("https://example.com/"))

    def test_scrape_site_deduplicates_repeated_links_on_same_page(self):
        html = """
        <html><body>
          <article><h2 class="entry-title"><a href="/story-1">A very specific unique headline about health policy</a></h2></article>
          <a href="/story-1">A very specific unique headline about health policy</a>
        </body></html>
        """
        with unittest.mock.patch("collectors.web_scraper.fetch_url", return_value=html):
            items = web_scraper.scrape_site({
                "name": "Test Site", "url": "https://example.com/", "language": "en",
                "category": "local_online", "link_selector": "h2.entry-title a",
            })
        self.assertEqual(len(items), 1)

    def test_scrape_site_sets_published_at_none(self):
        """No date is available from a listing page -- items should have
        published_at=None, which the scan pipeline treats as 'keep it'."""
        with unittest.mock.patch("collectors.web_scraper.fetch_url", return_value=self.WORDPRESS_STYLE_HTML):
            items = web_scraper.scrape_site({
                "name": "Test Site", "url": "https://example.com/", "language": "en",
                "category": "local_online", "link_selector": "h2.entry-title a",
            })
        self.assertTrue(all(i["published_at"] is None for i in items))

    def test_scrape_site_handles_network_failure_gracefully(self):
        with unittest.mock.patch("collectors.web_scraper.fetch_url", side_effect=Exception("connection refused")):
            items = web_scraper.scrape_site({
                "name": "Test Site", "url": "https://example.com/", "language": "en",
                "category": "local_online", "link_selector": None,
            })
        self.assertEqual(items, [])

    def test_scrape_site_handles_malformed_html_gracefully(self):
        with unittest.mock.patch("collectors.web_scraper.fetch_url", return_value="<html><body><article>"):
            items = web_scraper.scrape_site({
                "name": "Test Site", "url": "https://example.com/", "language": "en",
                "category": "local_online", "link_selector": None,
            })
        self.assertIsInstance(items, list)  # BeautifulSoup tolerates malformed HTML; should not crash

    def test_collect_aggregates_across_all_configured_sites(self):
        with unittest.mock.patch("collectors.web_scraper.fetch_url", return_value=self.WORDPRESS_STYLE_HTML):
            items = web_scraper.collect()
        # config.SCRAPE_SITES currently has 4 sites, each returning 2 items from this fixture
        self.assertEqual(len(items), len(config.SCRAPE_SITES) * 2)

    def test_items_include_correct_source_metadata(self):
        with unittest.mock.patch("collectors.web_scraper.fetch_url", return_value=self.WORDPRESS_STYLE_HTML):
            items = web_scraper.scrape_site({
                "name": "IGIHE", "url": "https://en.igihe.com/", "language": "en",
                "category": "local_online", "link_selector": "h2.entry-title a",
            })
        self.assertTrue(all(i["source_name"] == "IGIHE" for i in items))
        self.assertTrue(all(i["source_category"] == "local_online" for i in items))


if __name__ == "__main__":
    unittest.main()
