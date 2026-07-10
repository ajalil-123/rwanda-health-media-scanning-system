"""
Generates a Word document matching RBC's existing Health Sector Media
Review template, from whatever an editor has marked "included" in the
web app for a given scan.

Template structure being replicated (see the sample report):
    <Media Review | Day Date>
    <one-line intro>
    Highlights: (bulleted, title + linked source)
    I. Local Media
        A. Online
        B. Radio & TV        <- not yet automated, placeholder only
    II. International Media
    III. Research Findings
    IV. Social Media Trends  <- not yet automated, placeholder only

Entry format: "Title (Source): summary text. Readmore" with "Readmore"
hyperlinked to the item's URL, matching the sample report exactly.
"""

import os
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

import db

LINK_COLOR = RGBColor(0x05, 0x63, 0xC1)

# Resolve relative to this file's location (the project root), not the
# process's current working directory -- keeps output/ consistent whether
# this is invoked via scan.py, the web app, or a script launched from
# somewhere else entirely.
_PROJECT_ROOT = os.path.dirname(os.path.abspath(__file__))
_DEFAULT_OUTPUT_DIR = os.path.join(_PROJECT_ROOT, "output")


def _add_hyperlink(paragraph, url, text, bold=True):
    """python-docx has no built-in hyperlink support -- this builds the
    required OOXML by hand so 'Readmore' (and Highlight titles) render as
    real, clickable Word hyperlinks rather than plain blue text."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    if bold:
        b = OxmlElement("w:b")
        rpr.append(b)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rpr.append(underline)

    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def _entry_paragraph(doc, item):
    """One report entry: 'Title (Source): summary. Readmore' -- matching
    the sample report's exact style (bold title+source, plain summary,
    hyperlinked Readmore)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)

    title_run = p.add_run(f"{item['title']} ({item['source_name']}): ")
    title_run.bold = True

    summary = (item.get("editor_summary") or "").strip()
    if summary:
        p.add_run(summary + " ")

    _add_hyperlink(p, item["url"], "Readmore")
    return p


def _section_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading


def _placeholder_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"[{text}]")
    run.italic = True
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    return p


def _format_report_date(window_end):
    return window_end.strftime("%A %B %-d, %Y") if os.name != "nt" else window_end.strftime("%A %B %#d, %Y")


def generate_report(scan_id, db_path=None, output_dir=None):
    """
    Builds the Word report for one scan from its editor-included items and
    writes it to output_dir (defaults to <project root>/output). Returns
    the file path.

    Raises ValueError if the scan doesn't exist or has no included items
    yet (nothing to generate -- the editor needs to review first).
    """
    if output_dir is None:
        output_dir = _DEFAULT_OUTPUT_DIR
    scan_info = db.get_scan(scan_id, db_path=db_path)
    if not scan_info:
        raise ValueError(f"No scan found with id {scan_id}")

    items = db.get_included_items(scan_id, db_path=db_path)
    if not items:
        raise ValueError(
            f"Scan {scan_id} has no items marked 'included' yet. "
            "Review the shortlist in the web app and check 'Include in report' "
            "for at least one item before generating."
        )

    window_end = datetime.fromisoformat(scan_info["window_end"])
    if window_end.tzinfo is None:
        window_end = window_end.replace(tzinfo=timezone.utc)

    by_category = {"local_online": [], "international": [], "research": []}
    for item in items:
        by_category.setdefault(item["source_category"], []).append(item)

    doc = Document()

    # Title -----------------------------------------------------------------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    label = "Special Media Review" if scan_info["mode"] == "weekly" else "Media Review"
    run = title.add_run(f"{label} | {_format_report_date(window_end)}")
    run.bold = True
    run.font.size = Pt(14)

    intro = doc.add_paragraph()
    intro_text = (
        "This report covers weekly health related news in Rwanda."
        if scan_info["mode"] == "weekly"
        else "This report covers today's health related news in Rwanda."
    )
    intro_run = intro.add_run(intro_text)
    intro_run.italic = True

    # Highlights --------------------------------------------------------------
    _section_heading(doc, "Highlights:", level=1)
    highlights = sorted(items, key=lambda i: (i.get("highlight_score") or 0), reverse=True)[:10]
    for item in highlights:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{item['title']} ")
        run.bold = True
        run.italic = True
        _add_hyperlink(p, item["url"], f"({item['source_name']})", bold=True)

    # I. Local Media ------------------------------------------------------------
    _section_heading(doc, "I. Local Media", level=1)
    _section_heading(doc, "A. Online", level=2)
    if by_category["local_online"]:
        for item in by_category["local_online"]:
            _entry_paragraph(doc, item)
    else:
        _placeholder_note(doc, "No local online items were marked included for this period")

    _section_heading(doc, "B. Radio & TV", level=2)
    _placeholder_note(
        doc,
        "Not yet automated -- radio/TV (YouTube) collection is a planned future phase. "
        "Add items manually here, following the same 'Title (Source): summary. Watch' style."
    )

    # II. International Media -----------------------------------------------------
    _section_heading(doc, "II. International Media", level=1)
    if by_category["international"]:
        for item in by_category["international"]:
            _entry_paragraph(doc, item)
    else:
        _placeholder_note(doc, "No international items were marked included for this period")

    # III. Research Findings ---------------------------------------------------------
    _section_heading(doc, "III. Research Findings", level=1)
    if by_category["research"]:
        for item in by_category["research"]:
            _entry_paragraph(doc, item)
    else:
        _placeholder_note(doc, "No research items were marked included for this period")

    # IV. Social Media Trends ------------------------------------------------------------
    _section_heading(doc, "IV. Social Media Trends", level=1)
    _placeholder_note(
        doc,
        "Not yet automated -- social media (X) collection is a planned future phase. "
        "Add items manually here, following the same '@handle: summary. Readmore' style."
    )

    os.makedirs(output_dir, exist_ok=True)
    date_tag = window_end.strftime("%Y-%m-%d")
    out_path = os.path.join(output_dir, f"{scan_info['mode']}_report_{date_tag}.docx")
    doc.save(out_path)
    return out_path
